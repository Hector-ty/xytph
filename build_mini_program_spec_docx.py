import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\thxy")
SOURCE = ROOT / "mini_program_unified_spec.md"
OUTPUT_DIR = ROOT / "output"
OUTPUT = OUTPUT_DIR / "内蒙古大学校园低碳日小程序完整开发规格与AI一次性交付提示词_V2.0.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "5F6B76"
TABLE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
BORDER = "B8C4D1"
CODE_FILL = "F2F4F7"


def set_run_font(run, ascii_font="Calibri", east_asia_font="Microsoft YaHei",
                 size=None, color=None, bold=None, italic=None):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:cs"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, ascii_font="Calibri", east_asia_font="Microsoft YaHei",
                   size=None, color=None, bold=None):
    style.font.name = ascii_font
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:cs"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)


def set_paragraph_shading(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_border_left(paragraph, color=BLUE, size=14, space=8):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = pbdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        pbdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.find(qn("w:tcMar"))
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcmar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=6):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def calculate_column_widths(rows):
    count = max(len(row) for row in rows)
    if count == 1:
        return [CONTENT_WIDTH_DXA]
    if count == 2:
        first_max = max(len(row[0]) if row else 0 for row in rows)
        if first_max <= 28:
            return [2600, CONTENT_WIDTH_DXA - 2600]

    weights = []
    for index in range(count):
        lengths = []
        for row in rows:
            text = row[index] if index < len(row) else ""
            ascii_count = sum(1 for char in text if ord(char) < 128)
            lengths.append(max(2.0, len(text) - ascii_count * 0.45))
        value = max(lengths)
        weights.append(min(max(value, 7), 42))

    minimum = 920 if count >= 5 else 1100
    remaining = CONTENT_WIDTH_DXA - minimum * count
    if remaining <= 0:
        return [CONTENT_WIDTH_DXA // count] * (count - 1) + [
            CONTENT_WIDTH_DXA - (CONTENT_WIDTH_DXA // count) * (count - 1)
        ]
    total_weight = sum(weights)
    widths = [minimum + int(remaining * weight / total_weight) for weight in weights]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tblpr = tbl.tblPr

    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.insert(0, tblw)
    tblw.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tblind.set(qn("w:type"), "dxa")

    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.insert(0, tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    tbl_header = trpr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        trpr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_row_cant_split(row):
    trpr = row._tr.get_or_add_trPr()
    cant_split = trpr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        trpr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def create_numbering(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multilevel = OxmlElement("w:multiLevelType")
    multilevel.set(qn("w:val"), "singleLevel")
    abstract.append(multilevel)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    numfmt = OxmlElement("w:numFmt")
    numfmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(numfmt)
    lvltext = OxmlElement("w:lvlText")
    lvltext.set(qn("w:val"), "\u2022" if kind == "bullet" else "%1.")
    lvl.append(lvltext)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)

    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    lvl.append(ppr)

    if kind == "bullet":
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), "Symbol")
        rfonts.set(qn("w:hAnsi"), "Symbol")
        rpr.append(rfonts)
        lvl.append(rpr)

    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def assign_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.insert(0, numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, numid])


def add_inline_runs(paragraph, text, default_size=None, default_color=None,
                    default_bold=None, default_italic=None):
    token_pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    for part in token_pattern.split(text):
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(
                run,
                ascii_font="Consolas",
                east_asia_font="Microsoft YaHei",
                size=default_size or 10,
                color=DARK_BLUE,
                bold=False,
            )
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(
                run,
                size=default_size,
                color=default_color,
                bold=True,
                italic=default_italic,
            )
        else:
            run = paragraph.add_run(part)
            set_run_font(
                run,
                size=default_size,
                color=default_color,
                bold=default_bold,
                italic=default_italic,
            )


def add_metadata_paragraph(doc, line):
    paragraph = doc.add_paragraph(style="Metadata")
    if "：" in line:
        label, value = line.split("：", 1)
        label_run = paragraph.add_run(f"{label}：")
        set_run_font(label_run, size=10, color=NAVY, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10, color=MUTED)
    else:
        add_inline_runs(paragraph, line, default_size=10, default_color=MUTED)
    return paragraph


def add_callout(doc, text):
    paragraph = doc.add_paragraph(style="Callout")
    set_paragraph_shading(paragraph, CALLOUT_FILL)
    set_paragraph_border_left(paragraph)
    add_inline_runs(paragraph, text, default_size=10.5, default_color=DARK_BLUE)
    return paragraph


def add_code_block(doc, lines):
    paragraph = doc.add_paragraph(style="Code Block")
    set_paragraph_shading(paragraph, CODE_FILL)
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "3")
        node.set(qn("w:color"), "D9DEE5")
        pbdr.append(node)
    ppr.append(pbdr)

    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(
            run,
            ascii_font="Consolas",
            east_asia_font="Microsoft YaHei",
            size=8.5,
            color="263238",
        )
        if index < len(lines) - 1:
            run.add_break()
    return paragraph


def add_markdown_table(doc, rows):
    widths = calculate_column_widths(rows)
    table = doc.add_table(rows=len(rows), cols=len(widths))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for row_index, values in enumerate(rows):
        set_row_cant_split(table.rows[row_index])
        for col_index, cell in enumerate(table.rows[row_index].cells):
            value = values[col_index] if col_index < len(values) else ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, TABLE_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if len(value) <= 18 and row_index == 0
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            if row_index == 0:
                paragraph.paragraph_format.keep_with_next = True
            add_inline_runs(
                paragraph,
                value,
                default_size=9.2 if row_index else 9.5,
                default_color=NAVY if row_index == 0 else "202124",
                default_bold=(row_index == 0),
            )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.space_before = Pt(0)
    return table


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal, size=11, color="202124")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    set_style_font(title, size=24, color=NAVY, bold=True)
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.line_spacing = 1.05
    title_ppr = title.element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, size=12.5, color=MUTED, bold=False)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.line_spacing = 1.15

    h1 = styles["Heading 1"]
    set_style_font(h1, size=16, color=BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.1
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    set_style_font(h2, size=13, color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing = 1.1
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    set_style_font(h3, size=12, color=DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.line_spacing = 1.1
    h3.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        set_style_font(style, size=11, color="202124")
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    metadata = styles.add_style("Metadata", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(metadata, size=10, color=MUTED)
    metadata.paragraph_format.space_before = Pt(0)
    metadata.paragraph_format.space_after = Pt(2)
    metadata.paragraph_format.line_spacing = 1.1

    callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(callout, size=10.5, color=DARK_BLUE)
    callout.paragraph_format.left_indent = Inches(0.15)
    callout.paragraph_format.right_indent = Inches(0.1)
    callout.paragraph_format.space_before = Pt(8)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.line_spacing = 1.2

    code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(code, ascii_font="Consolas", size=8.5, color="263238")
    code.paragraph_format.left_indent = Inches(0.12)
    code.paragraph_format.right_indent = Inches(0.08)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.05


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("内蒙古大学校园低碳日小程序完整开发规格  |  V2.0 完整交付版")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    prefix = fp.add_run("2026-06-05  |  第 ")
    set_run_font(prefix, size=9, color=MUTED)
    add_page_field(fp)
    suffix = fp.add_run(" 页")
    set_run_font(suffix, size=9, color=MUTED)

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        line = lines[index].strip().strip("|")
        cells = [cell.strip() for cell in line.split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def build():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    configure_styles(doc)
    configure_document(doc)
    doc.core_properties.title = "内蒙古大学校园低碳日小程序完整开发规格与 AI 一次性交付提示词"
    doc.core_properties.subject = "校园低碳日小程序完整开发规格"
    doc.core_properties.keywords = "微信小程序, 低碳日, 开发规格, AI提示词"

    index = 0
    in_code = False
    code_lines = []
    current_list_kind = None
    current_num_id = None
    title_added = False
    metadata_prefixes = ("版本：", "日期：", "适用对象：", "系统范围：")

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if in_code:
            if stripped.startswith("```"):
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                code_lines.append(raw)
            index += 1
            continue

        if stripped.startswith("```"):
            current_list_kind = None
            current_num_id = None
            in_code = True
            code_lines = []
            index += 1
            continue

        if not stripped:
            current_list_kind = None
            current_num_id = None
            index += 1
            continue

        if stripped.startswith("|"):
            current_list_kind = None
            current_num_id = None
            rows, index = parse_table(lines, index)
            add_markdown_table(doc, rows)
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            current_list_kind = None
            current_num_id = None
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2)
            if level == 1 and not title_added:
                paragraph = doc.add_paragraph(style="Title")
                add_inline_runs(paragraph, heading_text, default_size=24, default_color=NAVY, default_bold=True)
                subtitle = doc.add_paragraph(style="Subtitle")
                add_inline_runs(
                    subtitle,
                    "学生端 / 工作人员端 / 管理后台 / 数据大屏 / 后端接口 / 数据模型",
                    default_size=12.5,
                    default_color=MUTED,
                )
                title_added = True
            else:
                style = {2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}.get(level, "Heading 1")
                paragraph = doc.add_paragraph(style=style)
                if heading_text.startswith("附录"):
                    paragraph.paragraph_format.page_break_before = True
                add_inline_runs(paragraph, heading_text)
            index += 1
            continue

        if stripped.startswith(">"):
            current_list_kind = None
            current_num_id = None
            add_callout(doc, stripped[1:].strip())
            index += 1
            continue

        bullet_match = re.match(r"^-\s+(.+)$", stripped)
        number_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet_match or number_match:
            kind = "bullet" if bullet_match else "number"
            text_value = bullet_match.group(1) if bullet_match else number_match.group(1)
            if current_list_kind != kind:
                current_list_kind = kind
                current_num_id = create_numbering(doc, kind)
            paragraph = doc.add_paragraph(style="List Bullet" if kind == "bullet" else "List Number")
            assign_numbering(paragraph, current_num_id)
            add_inline_runs(paragraph, text_value)
            index += 1
            continue

        current_list_kind = None
        current_num_id = None
        if stripped.startswith(metadata_prefixes):
            add_metadata_paragraph(doc, stripped)
        else:
            paragraph = doc.add_paragraph(style="Normal")
            add_inline_runs(paragraph, stripped)
        index += 1

    if in_code and code_lines:
        add_code_block(doc, code_lines)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
