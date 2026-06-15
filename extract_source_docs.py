from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(r"D:\thxy")
SOURCE_DIR = ROOT / "source_docs"
OUT_DIR = ROOT / "extracted"


def iter_block_items(doc):
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def clean(text):
    return " ".join(text.replace("\xa0", " ").split())


def extract_doc(path):
    doc = Document(path)
    lines = [f"# {path.stem}", ""]
    paragraph_count = 0
    table_count = 0

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = clean(block.text)
            if not text:
                continue
            paragraph_count += 1
            style = block.style.name if block.style else "Normal"
            lines.append(f"[P:{style}] {text}")
            lines.append("")
        else:
            table_count += 1
            lines.append(f"## [TABLE {table_count}]")
            for row in block.rows:
                cells = [clean(cell.text) for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
            lines.append("")

    for section_index, section in enumerate(doc.sections, start=1):
        header_text = " ".join(
            clean(p.text) for p in section.header.paragraphs if clean(p.text)
        )
        footer_text = " ".join(
            clean(p.text) for p in section.footer.paragraphs if clean(p.text)
        )
        if header_text:
            lines.extend([f"## [HEADER {section_index}]", header_text, ""])
        if footer_text:
            lines.extend([f"## [FOOTER {section_index}]", footer_text, ""])

    return "\n".join(lines), paragraph_count, table_count


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    index_lines = ["# Extraction index", ""]
    for path in sorted(SOURCE_DIR.glob("*.docx")):
        text, paragraphs, tables = extract_doc(path)
        output_path = OUT_DIR / f"{path.stem}.md"
        output_path.write_text(text, encoding="utf-8")
        index_lines.append(
            f"- {path.name}: {paragraphs} paragraphs, {tables} tables -> {output_path.name}"
        )
    (OUT_DIR / "_index.md").write_text("\n".join(index_lines), encoding="utf-8")
    print("\n".join(index_lines))


if __name__ == "__main__":
    main()
