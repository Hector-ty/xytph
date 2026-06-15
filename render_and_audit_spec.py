import json
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from pypdf import PdfReader


ROOT = Path(r"D:\thxy")
DOCX = ROOT / "output" / "内蒙古大学校园低碳日小程序统一开发规格与AI实施提示词_V1.0.docx"
PDF = ROOT / "rendered_spec_word" / "spec.pdf"
PNG_DIR = ROOT / "rendered_spec_word" / "pages"
REPORT = ROOT / "rendered_spec_word" / "audit.json"


def dxa_value(element, tag, attribute="w:w"):
    node = element.find(qn(tag))
    if node is None:
        return None
    value = node.get(qn(attribute))
    return int(value) if value is not None else None


def audit_docx():
    doc = Document(DOCX)
    table_issues = []
    for table_index, table in enumerate(doc.tables, start=1):
        tbl_pr = table._tbl.tblPr
        table_width = dxa_value(tbl_pr, "w:tblW")
        table_indent = dxa_value(tbl_pr, "w:tblInd")
        grid_widths = [
            int(col.get(qn("w:w"))) for col in table._tbl.tblGrid
            if col.get(qn("w:w")) is not None
        ]
        cell_width_issues = []
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                tc_width = dxa_value(cell._tc.get_or_add_tcPr(), "w:tcW")
                expected = grid_widths[min(cell_index - 1, len(grid_widths) - 1)]
                if tc_width != expected:
                    cell_width_issues.append(
                        {"row": row_index, "cell": cell_index, "actual": tc_width, "expected": expected}
                    )
        if table_width != 9360 or table_indent != 120 or sum(grid_widths) != 9360 or cell_width_issues:
            table_issues.append(
                {
                    "table": table_index,
                    "table_width": table_width,
                    "table_indent": table_indent,
                    "grid_sum": sum(grid_widths),
                    "cell_width_issues": cell_width_issues[:10],
                }
            )

    section = doc.sections[0]
    report = {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "page_width_inches": section.page_width.inches,
        "page_height_inches": section.page_height.inches,
        "margins_inches": {
            "top": section.top_margin.inches,
            "right": section.right_margin.inches,
            "bottom": section.bottom_margin.inches,
            "left": section.left_margin.inches,
        },
        "table_geometry_issues": table_issues,
        "required_text": {},
    }

    all_text = "\n".join(p.text for p in doc.paragraphs)
    for phrase in (
        "锁定业务规则",
        "API 统一契约",
        "数据模型与约束",
        "可直接发送给 AI 的主提示词",
        "测试与验收清单",
    ):
        report["required_text"][phrase] = phrase in all_text

    with ZipFile(DOCX) as archive:
        report["contains_numbering_xml"] = "word/numbering.xml" in archive.namelist()
        report["contains_header_xml"] = any(name.startswith("word/header") for name in archive.namelist())
        report["contains_footer_xml"] = any(name.startswith("word/footer") for name in archive.namelist())
    return report


def render_pdf():
    return len(PdfReader(PDF).pages)


def main():
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    report = audit_docx()
    report["pdf_pages"] = render_pdf()
    report["rendered_png_pages"] = len(list(PNG_DIR.glob("page-*.png")))
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
